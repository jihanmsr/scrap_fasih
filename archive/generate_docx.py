import os
import subprocess
import sys

def install_and_import():
    try:
        import docx
        # Make sure it's the right docx
        if not hasattr(docx, 'Document'):
            raise ImportError
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "uninstall", "-y", "docx"])
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

install_and_import()
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Dokumentasi Perjalanan & Panduan Aplikasi "Monitoring Sensus Ekonomi 2026"', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Dokumen ini berisi dokumentasi lengkap dari awal mula pengembangan aplikasi Dashboard Monitoring Sensus Ekonomi 2026 BPS Provinsi Sulawesi Tengah, penjelasan fungsionalitas setiap halaman, sumber pengambilan data, hingga fitur-fitur yang tersedia di dalam aplikasi.")
    
    # 1. Latar Belakang
    doc.add_heading('1. Latar Belakang & Perjalanan Aplikasi', level=1)
    doc.add_paragraph("Aplikasi ini dibangun untuk memfasilitasi kebutuhan Badan Pusat Statistik (BPS) Provinsi Sulawesi Tengah dalam memantau secara real-time progres pelaksanaan Sensus Ekonomi 2026.")
    doc.add_paragraph("Pada awalnya, aplikasi ini berfokus pada Pemantauan Email Usaha Besar (UB) untuk melacak status kuesioner yang dikirimkan via email (apakah terkirim, bounced, dibuka, atau diklik). Seiring berjalannya waktu, aplikasi ini berkembang menjadi sebuah Dashboard Monitoring terpadu yang juga mencakup pemantauan progres pencacahan Sensus Ekonomi Umum, Sensus Ekonomi Usaha Besar, dan Alokasi Penugasan Petugas.")
    
    # 2. Sumber Pengambilan Data
    doc.add_heading('2. Sumber Pengambilan Data (Data Sources)', level=1)
    doc.add_paragraph("Aplikasi ini mengandalkan beberapa skrip Python di backend (berjalan di balik layar atau via cron/task scheduler) untuk menarik data dari server FASIH dan sistem email, kemudian menyimpannya ke dalam file lokal (.js, .json, .csv) yang akan dibaca oleh halaman aplikasi (index.html).")
    
    sources = [
        ("scrape_via_api.py / scrape_sync.py", "Mengambil data progres capaian sensus secara real-time dari API FASIH BPS. Melakukan autentikasi menggunakan bearer token, kemudian menarik data agregat dan rincian per kabupaten/kota. Hasilnya disimpan ke dalam file seperti data.js atau sync_data.js."),
        ("scrape_assign.py", "Mengambil data alokasi penugasan petugas (siapa petugas yang ditugaskan, berapa beban kerjanya, mana yang belum dialokasikan). Menyimpannya ke assign_data.js."),
        ("generate_ipas_report.py", "Menghitung persentase capaian, kalkulasi progres harian (kenaikan hari ini vs kemarin), dan rekapitulasi 'Sisa Usaha'. Mengolah raw data menjadi format tabel dan grafik."),
        ("Sistem Pemantauan Email (Supabase / CSV)", "Menyimpan log pengiriman email (all_email_history.csv dan bounced_emails.xlsx). Memantau status setiap perusahaan (Bounced, Delivered, Opened, Clicked).")
    ]
    
    for title_text, desc in sources:
        p = doc.add_paragraph(style='List Bullet')
        runner = p.add_run(title_text + ": ")
        runner.bold = True
        p.add_run(desc)

    # 3. Fitur Utama
    doc.add_heading('3. Fitur Utama & Tampilan Antarmuka (UI/UX)', level=1)
    doc.add_paragraph("Aplikasi ini dirancang dengan gaya desain modern, menggunakan glassmorphism, animasi interaktif (micro-animations), dan mendukung 2 Mode Tampilan (Light Mode & Dark Mode).")
    
    note = doc.add_paragraph()
    note_run = note.add_run("Catatan: Pengguna dapat mengklik ikon Matahari/Bulan di sudut kiri bawah sidebar untuk beralih antara tampilan Terang dan Gelap.")
    note_run.italic = True
    
    # A. SE Umum
    doc.add_heading('A. Halaman Sensus Ekonomi Umum', level=2)
    doc.add_paragraph("Halaman ini difokuskan pada pemantauan target usaha secara umum (Prelist).")
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run("Menampilkan Total Target Usaha yang harus disensus.")
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run("Membagi status dokumen menjadi: Draft (sedang dikerjakan), Open (baru dibuka), dan Submitted (selesai).")
    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run("Menampilkan grafik progres capaian dan rincian tabel per Kabupaten/Kota.")
    
    ss_p1 = doc.add_paragraph()
    ss_run1 = ss_p1.add_run("[ TEMPATKAN SCREENSHOT SE UMUM - LIGHT MODE DI SINI ]\n[ TEMPATKAN SCREENSHOT SE UMUM - DARK MODE DI SINI ]")
    ss_run1.font.color.rgb = RGBColor(255, 0, 0)
    ss_run1.bold = True
    ss_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # B. SE UB
    doc.add_heading('B. Halaman Sensus Ekonomi Usaha Besar (UB)', level=2)
    doc.add_paragraph("Mirip dengan Sensus Ekonomi Umum, namun dikhususkan untuk memantau progres pendataan perusahaan dalam kategori Usaha Besar.")
    
    ss_p2 = doc.add_paragraph()
    ss_run2 = ss_p2.add_run("[ TEMPATKAN SCREENSHOT SE UB - LIGHT MODE DI SINI ]\n[ TEMPATKAN SCREENSHOT SE UB - DARK MODE DI SINI ]")
    ss_run2.font.color.rgb = RGBColor(255, 0, 0)
    ss_run2.bold = True
    ss_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # C. Email
    doc.add_heading('C. Halaman Pemantauan Email Usaha Besar', level=2)
    doc.add_paragraph("Halaman khusus yang memantau interaksi email blast kuesioner ke perusahaan Usaha Besar.")
    pe1 = doc.add_paragraph(style='List Bullet')
    pe1.add_run("Menampilkan statistik: Total Perusahaan, Bounced Emails (Gagal), Permanent Fail, Link Clicked, dan Opened Emails.")
    pe2 = doc.add_paragraph(style='List Bullet')
    pe2.add_run("Terdapat fitur Pencarian dan Filter berdasarkan Kabupaten/Kota dan Status, serta tampilan Tabel dan Kartu (Card).")
    
    ss_p3 = doc.add_paragraph()
    ss_run3 = ss_p3.add_run("[ TEMPATKAN SCREENSHOT EMAIL - LIGHT MODE DI SINI ]\n[ TEMPATKAN SCREENSHOT EMAIL - DARK MODE DI SINI ]")
    ss_run3.font.color.rgb = RGBColor(255, 0, 0)
    ss_run3.bold = True
    ss_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # D. Alokasi
    doc.add_heading('D. Halaman Alokasi Penugasan Petugas', level=2)
    doc.add_paragraph("Halaman manajerial untuk melihat distribusi beban kerja para petugas sensus.")
    pa1 = doc.add_paragraph(style='List Bullet')
    pa1.add_run("Memantau jumlah sampel/usaha yang sudah Dialokasikan (Assigned) dan Belum Dialokasikan (Unassigned).")
    pa2 = doc.add_paragraph(style='List Bullet')
    pa2.add_run("Terdapat sub-tab untuk memisahkan kategori Usaha Umum dan Usaha Besar (UB) beserta grafik donat visualisasi beban tugas.")

    ss_p4 = doc.add_paragraph()
    ss_run4 = ss_p4.add_run("[ TEMPATKAN SCREENSHOT ALOKASI - LIGHT MODE DI SINI ]\n[ TEMPATKAN SCREENSHOT ALOKASI - DARK MODE DI SINI ]")
    ss_run4.font.color.rgb = RGBColor(255, 0, 0)
    ss_run4.bold = True
    ss_p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 4. Keunggulan
    doc.add_heading('4. Keunggulan & Teknologi', level=1)
    doc.add_paragraph("1. Responsif & Real-time: Dashboard menggunakan HTML, CSS (Vanilla), dan JavaScript dengan Chart.js.")
    doc.add_paragraph("2. Auto-Refresh: Terdapat fitur refresh otomatis untuk memastikan data terbaru.")
    doc.add_paragraph("3. Penyimpanan Lokal & Backup: Data scraping disimpan dalam format JSON/CSV/JS sebagai cache lokal sehingga loading sangat cepat.")
    
    doc.save('Dokumentasi_Aplikasi_Scrap_Fasih.docx')
    print("Document saved as Dokumentasi_Aplikasi_Scrap_Fasih.docx")

if __name__ == '__main__':
    create_doc()
