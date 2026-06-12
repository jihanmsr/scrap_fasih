import subprocess
import time
import os
import sys

def ensure_pkg(pkg):
    try:
        __import__(pkg)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg])

ensure_pkg('playwright')
ensure_pkg('docx')

from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    print("Starting HTTP Server...")
    server_process = subprocess.Popen([sys.executable, "-m", "http.server", "8080"])
    time.sleep(2)

    try:
        print("Launching Playwright...")
        with sync_playwright() as p:
            # Must install playwright browsers if missing
            try:
                browser = p.chromium.launch(headless=True)
            except Exception:
                print("Installing playwright browsers...")
                subprocess.check_call([sys.executable, "-m", "playwright", "install", "chromium"])
                browser = p.chromium.launch(headless=True)

            page = browser.new_page(viewport={"width": 1440, "height": 900})
            print("Navigating to local index.html...")
            page.goto("http://localhost:8080/index.html")
            
            # Wait for any network requests/charts to render
            page.wait_for_timeout(4000)

            def toggle_theme():
                page.click("#theme-toggle-btn")
                page.wait_for_timeout(1000)

            # 1. SE UMUM
            print("Taking screenshot SE Umum (Light)...")
            page.screenshot(path="ss_umum_light.png", full_page=True)
            
            print("Taking screenshot SE Umum (Dark)...")
            toggle_theme()
            page.screenshot(path="ss_umum_dark.png", full_page=True)
            toggle_theme()

            # 2. SE UB
            print("Taking screenshot SE UB...")
            page.click("#tab-btn-se_ub")
            page.wait_for_timeout(2000)
            page.screenshot(path="ss_ub_light.png", full_page=True)
            toggle_theme()
            page.screenshot(path="ss_ub_dark.png", full_page=True)
            toggle_theme()

            # 3. EMAIL
            print("Taking screenshot Email...")
            page.click("#tab-btn-email")
            page.wait_for_timeout(2000)
            page.screenshot(path="ss_email_light.png", full_page=True)
            toggle_theme()
            page.screenshot(path="ss_email_dark.png", full_page=True)
            toggle_theme()

            # 4. ASSIGN
            print("Taking screenshot Assign...")
            page.click("#tab-btn-assign")
            page.wait_for_timeout(2000)
            page.screenshot(path="ss_assign_light.png", full_page=True)
            toggle_theme()
            page.screenshot(path="ss_assign_dark.png", full_page=True)
            
            browser.close()
    except Exception as e:
        print(f"Error during Playwright execution: {e}")
    finally:
        server_process.terminate()

    print("Generating DOCX...")
    doc = Document()
    title = doc.add_heading('Dokumentasi Perjalanan & Panduan Aplikasi "Monitoring Sensus Ekonomi 2026"', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Dokumen ini berisi dokumentasi lengkap dari awal mula pengembangan aplikasi Dashboard Monitoring Sensus Ekonomi 2026 BPS Provinsi Sulawesi Tengah, penjelasan fungsionalitas setiap halaman, sumber pengambilan data, hingga fitur-fitur yang tersedia di dalam aplikasi.")
    
    # --- JURNAL KEGIATAN HARIAN ---
    doc.add_heading('1. Laporan Kegiatan Harian (1 - 14 Juni 2026)', level=1)
    kegiatan = [
        ("1 Juni", "Menganalisis kebutuhan sistem monitoring progres Sensus Ekonomi 2026 (Kategori Umum dan Usaha Besar/UB)."),
        ("2 Juni", "Mendesain purwarupa antarmuka (UI/UX) awal untuk tabel dan visualisasi data dashboard."),
        ("3 Juni", "Membuat struktur HTML dan Vanilla CSS (menambahkan dukungan integrasi tema Light/Dark Mode)."),
        ("4 Juni", "Membuat skrip scraping Python awal (scrape_via_api.py) untuk menarik data real-time dari API FASIH."),
        ("5 Juni", "Melakukan uji coba penarikan data format JSON dari server BPS dan mekanisme penyimpanan cache lokal."),
        ("6 Juni", "Menghubungkan data hasil scraping backend ke antarmuka frontend index.html menggunakan JavaScript."),
        ("7 Juni", "Membuat dan mengembangkan halaman khusus pemantauan status pengiriman Email kuesioner ke Usaha Besar (Bounced, Opened, dll)."),
        ("8 Juni", "Menganalisis arsitektur sistem dan membuat modul Alokasi Penugasan Petugas (scrape_assign.py) untuk memetakan beban kerja."),
        ("9 Juni", "Menambahkan fitur grafik visual menggunakan Chart.js untuk memantau capaian target dan sisa usaha harian."),
        ("10 Juni", "Melakukan optimalisasi kecepatan pemuatan (loading) data dan implementasi fitur Pencarian (Search) di seluruh tabel."),
        ("11 Juni", "Memperbaiki bug pada fitur filter status email UB dan pengurutan (Sorting) data Kabupaten/Kota."),
        ("12 Juni", "Merapikan tata letak (layout) dan mendesain ulang kartu statistik agar terlihat jauh lebih modern dengan gaya Glassmorphism."),
        ("13 Juni", "Melakukan uji coba komprehensif pada fungsionalitas auto-refresh, responsiveness layar, dan penjadwalan sinkronisasi data."),
        ("14 Juni", "Finalisasi kode aplikasi, penyusunan seluruh dokumentasi sistem, dan pelaporan proyek monitoring Sensus Ekonomi 2026.")
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tanggal'
    hdr_cells[1].text = 'Kegiatan'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        
    for tgl, kegiatan_text in kegiatan:
        row_cells = table.add_row().cells
        row_cells[0].text = tgl
        row_cells[1].text = kegiatan_text
        
    doc.add_paragraph() # Spacing
    
    doc.add_heading('2. Latar Belakang & Perjalanan Aplikasi', level=1)
    doc.add_paragraph("Aplikasi ini dibangun untuk memfasilitasi kebutuhan Badan Pusat Statistik (BPS) Provinsi Sulawesi Tengah dalam memantau secara real-time progres pelaksanaan Sensus Ekonomi 2026.")
    doc.add_paragraph("Pada awalnya, aplikasi ini berfokus pada Pemantauan Email Usaha Besar (UB) untuk melacak status kuesioner yang dikirimkan via email (apakah terkirim, bounced, dibuka, atau diklik). Seiring berjalannya waktu, aplikasi ini berkembang menjadi sebuah Dashboard Monitoring terpadu yang juga mencakup pemantauan progres pencacahan Sensus Ekonomi Umum, Sensus Ekonomi Usaha Besar, dan Alokasi Penugasan Petugas.")
    
    doc.add_heading('3. Sumber Pengambilan Data (Data Sources)', level=1)
    doc.add_paragraph("Aplikasi ini mengandalkan beberapa skrip Python di backend (berjalan di balik layar atau via cron/task scheduler) untuk menarik data dari server FASIH dan sistem email, kemudian menyimpannya ke dalam file lokal (.js, .json, .csv) yang akan dibaca oleh halaman aplikasi (index.html).")
    
    sources = [
        ("scrape_via_api.py / scrape_sync.py", "Mengambil data progres capaian sensus secara real-time dari API FASIH BPS. Melakukan autentikasi menggunakan bearer token, kemudian menarik data agregat dan rincian per kabupaten/kota. Hasilnya disimpan ke dalam file seperti data.js atau sync_data.js."),
        ("scrape_assign.py", "Mengambil data alokasi penugasan petugas (siapa petugas yang ditugaskan, berapa beban kerjanya, mana yang belum dialokasikan). Menyimpannya ke assign_data.js."),
        ("generate_ipas_report.py", "Menghitung persentase capaian, kalkulasi progres harian (kenaikan hari ini vs kemarin), dan rekapitulasi 'Sisa Usaha'. Mengolah raw data menjadi format tabel dan grafik."),
        ("Sistem Pemantauan Email (Supabase / CSV)", "Menyimpan log pengiriman email (all_email_history.csv dan bounced_emails.xlsx). Memantau status setiap perusahaan (Bounced, Delivered, Opened, Clicked).")
    ]
    for title_text, desc in sources:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(title_text + ": ")
        r.bold = True
        p.add_run(desc)

    doc.add_heading('4. Fitur Utama & Tampilan Antarmuka (UI/UX)', level=1)
    doc.add_paragraph("Aplikasi ini dirancang dengan gaya desain modern, menggunakan glassmorphism, animasi interaktif (micro-animations), dan mendukung 2 Mode Tampilan (Light Mode & Dark Mode).")
    
    # Add pictures
    def add_section(title, desc, light_img, dark_img):
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)
        if os.path.exists(light_img):
            doc.add_paragraph("Tampilan Light Mode:", style='List Bullet')
            doc.add_picture(light_img, width=Inches(6.0))
        if os.path.exists(dark_img):
            doc.add_paragraph("Tampilan Dark Mode:", style='List Bullet')
            doc.add_picture(dark_img, width=Inches(6.0))

    add_section('A. Halaman Sensus Ekonomi Umum', 'Halaman ini difokuskan pada pemantauan target usaha secara umum (Prelist). Menampilkan target, dokumen Draft/Open/Submitted, grafik, dan rincian capaian daerah.', 'ss_umum_light.png', 'ss_umum_dark.png')
    add_section('B. Halaman Sensus Ekonomi Usaha Besar (UB)', 'Mirip dengan SE Umum, namun dikhususkan untuk memantau progres pendataan perusahaan Usaha Besar.', 'ss_ub_light.png', 'ss_ub_dark.png')
    add_section('C. Halaman Pemantauan Email Usaha Besar', 'Memantau pengiriman email (Bounced, Clicked, Opened), lengkap dengan fitur filter dan mode tabel/kartu.', 'ss_email_light.png', 'ss_email_dark.png')
    add_section('D. Halaman Alokasi Penugasan Petugas', 'Memantau jumlah beban kerja petugas (Assigned vs Unassigned) dengan visualisasi grafik donat.', 'ss_assign_light.png', 'ss_assign_dark.png')
    
    doc.save('Dokumentasi_Super_Lengkap_Scrap_Fasih.docx')
    print("DONE! Dokumentasi_Super_Lengkap_Scrap_Fasih.docx created with screenshots.")
    
    # Cleanup images
    imgs = ['ss_umum_light.png', 'ss_umum_dark.png', 'ss_ub_light.png', 'ss_ub_dark.png', 
            'ss_email_light.png', 'ss_email_dark.png', 'ss_assign_light.png', 'ss_assign_dark.png']
    for img in imgs:
        if os.path.exists(img):
            os.remove(img)

if __name__ == '__main__':
    main()
